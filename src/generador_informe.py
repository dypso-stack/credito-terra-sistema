"""
Módulo: Generador de Informe Word
Crédito Terra - Banco Guayaquil

Genera el informe técnico en formato Word (.docx) para cada evaluación,
listo para presentar al Comité de Crédito.
"""

import sys
import os
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    DOCX_DISPONIBLE = True
except ImportError:
    DOCX_DISPONIBLE = False


# Colores corporativos Banco Guayaquil
COLOR_MAGENTA = RGBColor(0xD4, 0x00, 0x7A) if DOCX_DISPONIBLE else None
COLOR_OSCURO = RGBColor(0x1A, 0x1A, 0x2E) if DOCX_DISPONIBLE else None
COLOR_VERDE = RGBColor(0x00, 0x7A, 0x33) if DOCX_DISPONIBLE else None
COLOR_ROJO = RGBColor(0xCC, 0x00, 0x00) if DOCX_DISPONIBLE else None
COLOR_GRIS = RGBColor(0x66, 0x66, 0x66) if DOCX_DISPONIBLE else None


def agregar_titulo(doc, texto, nivel=1, color=None, centrado=False):
    """Agrega un título con formato."""
    p = doc.add_heading(texto, level=nivel)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    if centrado:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def agregar_parrafo(doc, texto, negrita=False, color=None, tamanio=10):
    """Agrega un párrafo con formato."""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrita
    run.font.size = Pt(tamanio)
    if color:
        run.font.color.rgb = color
    return p


def agregar_tabla_datos(doc, datos: list, encabezados: list):
    """
    Agrega una tabla con encabezados y datos.
    datos: lista de listas [[fila1col1, fila1col2], ...]
    """
    tabla = doc.add_table(rows=1, cols=len(encabezados))
    tabla.style = 'Table Grid'

    # Encabezados
    fila_enc = tabla.rows[0].cells
    for i, enc in enumerate(encabezados):
        fila_enc[i].text = enc
        fila_enc[i].paragraphs[0].runs[0].bold = True
        fila_enc[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        fila_enc[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Datos
    for fila in datos:
        row = tabla.add_row().cells
        for i, valor in enumerate(fila):
            row[i].text = str(valor)
            row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    return tabla


def generar_informe_word(caso: dict, dictamen: dict, ruta_salida: str) -> str:
    """
    Genera el informe técnico Word para un caso evaluado.

    Parámetros:
        caso: Datos del cliente y proyecto
        dictamen: Resultado del motor de elegibilidad
        ruta_salida: Ruta donde guardar el .docx

    Retorna:
        Ruta del archivo generado
    """
    if not DOCX_DISPONIBLE:
        raise ImportError("python-docx no está instalado. Ejecuta: pip install python-docx")

    doc = Document()

    # --- Configurar márgenes ---
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


    #________________________________________
    # ENCABEZADO
    #________________________________________    

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BANCO GUAYAQUIL")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_MAGENTA

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Área de Sostenibilidad")
    run2.font.size = Pt(11)
    run2.font.color.rgb = COLOR_OSCURO

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("INFORME DE EVALUACIÓN DE ELEGIBILIDAD DE CRÉDITO TERRA")
    run3.bold = True
    run3.font.size = Pt(13)

    doc.add_paragraph()
   
   #________________________________________
   # SECCIÓN 1: INFORMACIÓN GENERAL
   #________________________________________

    agregar_titulo(doc, "1. Información general del caso", nivel=2, color=COLOR_MAGENTA)

    datos_generales = [
        ["\nID Caso", caso.get('id_caso', 'N/D')],
        ["Fecha de evaluación", dictamen.get('fecha_evaluacion', datetime.today().strftime('%d/%m/%Y'))],
        ["Responsable del levantamiento", caso.get('responsable', 'N/D')],
        ["Asesor comercial", caso.get('asesor_comercial', 'N/D')],
    ]
    for campo, valor in datos_generales:
        p = doc.add_paragraph()
        p.add_run(f"{campo}: ").bold = True
        p.add_run(str(valor) if valor else 'No declarado')

    doc.add_paragraph()

    #________________________________________
    # SECCIÓN 2: DATOS DEL CLIENTE
    #________________________________________

    agregar_titulo(doc, "2. Datos del cliente", nivel=2, color=COLOR_MAGENTA)

    datos_cliente = [
        ["\nRazón social", caso.get('razon_social', 'N/D')],
        ["RUC / Cédula", caso.get('ruc', 'N/D')],
        ["Segmento", caso.get('segmento', 'N/D')],
        ["Sector BG", caso.get('sector_bg', 'N/D')],
        ["Actividad principal", caso.get('actividad_principal', 'N/D')],
        ["Correo de contacto", caso.get('correo_contacto', 'N/D')],
    ]
    for campo, valor in datos_cliente:
        p = doc.add_paragraph()
        run = p.add_run(f"{campo}: ")
        run.bold = True
        p.add_run(str(valor) if valor else 'No declarado')

    doc.add_paragraph()

    #________________________________________
    # SECCIÓN 3: DATOS DEL CRÉDITO Y PROYECTO
    #________________________________________

    agregar_titulo(doc, "3. Datos del crédito y proyecto", nivel=2, color=COLOR_MAGENTA)

    monto = caso.get('monto_credito_usd', 0)
    monto_str = f"${monto:,.2f}" if isinstance(monto, (int, float)) else 'N/D'
    capex = caso.get('capex_usd', 0)
    capex_str = f"${capex:,.2f}" if isinstance(capex, (int, float)) else 'N/D'

    datos_credito = [
        ["\nDestino de Crédito Terra", caso.get('destino_terra', 'N/D')],
        ["Código Crédito Terra", caso.get('codigo_terra', 'N/D')],
        ["Monto solicitado", monto_str],
        ["CAPEX total del proyecto", capex_str],
        ["Etapa del proyecto", caso.get('etapa', 'N/D')],
        ["Tipo de sistema", caso.get('tipo_sistema', 'N/D')],
        ["Dirección del proyecto", caso.get('direccion_proyecto', 'N/D')],
    ]
    for campo, valor in datos_credito:
        p = doc.add_paragraph()
        run = p.add_run(f"{campo}: ")
        run.bold = True
        p.add_run(str(valor) if valor else 'No declarado')

    doc.add_paragraph()

    #________________________________________
    # SECCIÓN 4: INDICADORES AMBIENTALES
    #________________________________________
    
    agregar_titulo(doc, "4. Indicadores Ambientales", nivel=2, color=COLOR_MAGENTA)

    ind = dictamen.get('indicadores', {})
    tabla_ind = agregar_tabla_datos(
        doc,
        datos=[
            ["Energía renovable producida", f"{ind.get('energia_producida_mwh_anio', 0)} MWh/año"],
            ["Capacidad instalada", f"{ind.get('capacidad_kwp', 0)} kWp"],
            ["Cobertura energética", f"{ind.get('cobertura_energetica_pct', 0)}%"],
            ["Factor de emisión SNI", f"{ind.get('factor_emision_sni', 0.4567)} tCO₂/MWh"],
            ["GEI evitadas por año", f"{ind.get('gei_evitadas_anual_tco2', 0)} tCO₂/año"],
            ["GEI evitadas totales (vida útil)", f"{ind.get('gei_evitadas_total_tco2', 0)} tCO₂"],
            ["Vida útil del sistema", f"{ind.get('vida_util_anios', 0)} años"],
        ],
        encabezados=["Indicador", "Valor"]
    )

    doc.add_paragraph()

    #________________________________________
    # SECCIÓN 5: EVALUACIÓN DE CRITERIOS
    #________________________________________

    agregar_titulo(doc, "5. Evaluación de Criterios de Elegibilidad", nivel=2, color=COLOR_MAGENTA)

    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run(
        f"\nCriterios cumplidos: {dictamen.get('n_criterios_cumplidos', 0)} / "
        f"{dictamen.get('n_criterios_total', 10)}"
    )
    run_sub.bold = True

    doc.add_paragraph()

    for linea in dictamen.get('detalle_criterios', []):
        p = doc.add_paragraph(style='List Bullet')
        cumple = linea.startswith('✅')
        run = p.add_run(linea)
        run.font.color.rgb = COLOR_VERDE if cumple else COLOR_ROJO
        run.font.size = Pt(9)

    doc.add_paragraph()

    #________________________________________
    # SECCIÓN 6: DICTAMEN FINAL
    #________________________________________

    agregar_titulo(doc, "6. Dictamen final", nivel=2, color=COLOR_MAGENTA)

    dictamen_valor = dictamen.get('dictamen', 'N/D')
    color_dictamen = COLOR_VERDE if dictamen_valor == 'CUMPLE' else COLOR_ROJO

    p_dict = doc.add_paragraph()
    p_dict.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_dict = p_dict.add_run(f"DICTAMEN: {dictamen_valor}")
    run_dict.bold = True
    run_dict.font.size = Pt(14)
    run_dict.font.color.rgb = color_dictamen

    doc.add_paragraph()

    p_accion = doc.add_paragraph()
    run_ac = p_accion.add_run("Acción requerida: ")
    run_ac.bold = True
    p_accion.add_run(dictamen.get('accion_requerida', 'N/D'))

    doc.add_paragraph()

    #________________________________________
    # SECCIÓN 7: OBSERVACIONES Y PRÓXIMOS PASOS
    #________________________________________
   
    agregar_titulo(doc, "7. Observaciones y próximos pasos", nivel=2, color=COLOR_MAGENTA)

    p_obs = doc.add_paragraph()
    run_obs = p_obs.add_run("\nObservaciones: ")
    run_obs.bold = True
    p_obs.add_run(caso.get('observaciones', 'Sin observaciones registradas.'))

    p_prox = doc.add_paragraph()
    run_prox = p_prox.add_run("Próximos pasos: ")
    run_prox.bold = True
    p_prox.add_run(caso.get('proximos_pasos', 'Pendiente de definir.'))

    doc.add_paragraph()

    #________________________________________
    # PIE DE PÁGINA
    #________________________________________
    
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pie = p_pie.add_run(
        f"Documento generado automáticamente por el Sistema de Evaluación Crédito Terra\n"
        f"Banco Guayaquil | {datetime.today().strftime('%d/%m/%Y %H:%M')}"
    )
    run_pie.font.size = Pt(8)
    run_pie.font.color.rgb = COLOR_GRIS

    # Guardar
    Path(ruta_salida).parent.mkdir(parents=True, exist_ok=True)
    doc.save(ruta_salida)
    return ruta_salida


if __name__ == "__main__":
    # Prueba con datos de ejemplo
    sys.path.append(os.path.dirname(__file__))
    from motor_elegibilidad import MotorElegibilidad

    caso_demo = {
        'id_caso': 'CT-0001',
        'razon_social': 'Empresa Demo S.A.',
        'ruc': '0912345678001',
        'segmento': 'Empresarial',
        'sector_bg': 'INDUSTRIA ALIMENTICIA',
        'actividad_principal': 'Producción de alimentos procesados',
        'correo_contacto': 'demo@empresa.com',
        'responsable': 'Ana Torres',
        'asesor_comercial': 'Pedro Álvarez',
        'destino_terra': 'Energía Renovable',
        'codigo_terra': '04 ER. 1',
        'monto_credito_usd': 850000,
        'capex_usd': 1200000,
        'etapa': 'Factibilidad',
        'tipo_sistema': 'Conectado a la red (SFCR)',
        'direccion_proyecto': 'Km 15 vía Daule, Guayas',
        'energia_producida_mwh_anio': 2500,
        'capacidad_instalada_mwp': 1.8,
        'consumo_cliente_mwh_anio': 3200,
        'vida_util_anios': 25,
        'factor_emision_sni': 0.4567,
        'gei_evitadas_anual_tco2': 1141.75,
        'gei_evitadas_total_tco2': 28543.75,
        'uso_exclusivo_fondos': True,
        'evidencia_objetiva_proyecto': True,
        'indicadores_ambientales_completos': True,
        'cumple_normativa_ambiental': True,
        'no_categoria_a_ifc': True,
        'no_involucra_combustibles_fosiles': True,
        'no_impacto_negativo_otros_objetivos': True,
        'no_lista_exclusion_saras': True,
        'no_controversias_socioambientales': True,
        'no_es_credito_consumo': True,
        'observaciones': 'Proyecto viable con alto impacto ambiental.',
        'proximos_pasos': 'Proceder a comité de aprobación.',
    }

    motor = MotorElegibilidad(caso_demo)
    dictamen = motor.emitir_dictamen()

    ruta = "data/outputs/Informe_CT-0001_demo.docx"
    generar_informe_word(caso_demo, dictamen, ruta)
    print(f"✅ Informe generado: {ruta}")
